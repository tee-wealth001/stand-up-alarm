# creating a doxc file with the estimated budget for the Holistic Student Insights Dashboard project.

from datetime import datetime
from docx import Document

# Create the Word Document
doc = Document()
doc.add_heading("Estimated Development Budget: Holistic Student Insights Dashboard", 0)

# Add date
doc.add_paragraph(f'Date: {datetime.today().strftime("%B %d, %Y")}\n')

# Development Costs
doc.add_heading("1. Development Costs (MVP)", level=1)
dev_table = doc.add_table(rows=1, cols=4)
hdr_cells = dev_table.rows[0].cells
hdr_cells[0].text = "Role"
hdr_cells[1].text = "Weekly Rate"
hdr_cells[2].text = "Duration"
hdr_cells[3].text = "Total"

roles = [
    ("Frontend Developer", "£1,200", "10 weeks", "£12,000"),
    ("Backend Developer", "£1,300", "10 weeks", "£13,000"),
    ("Data/AI Engineer", "£1,500", "4 weeks", "£6,000"),
    ("Tech Lead / Architect", "£1,600", "4 weeks (part-time)", "£6,400"),
    ("Project Manager", "£1,000", "10 weeks", "£10,000"),
    ("QA/Test Engineer", "£1,000", "4 weeks", "£4,000"),
]
for role in roles:
    row_cells = dev_table.add_row().cells
    for i in range(4):
        row_cells[i].text = role[i]

doc.add_paragraph("Estimated Development Team Cost (MVP): £45,000 – £55,000")

# Infrastructure Costs
doc.add_heading("2. Infrastructure & Hosting (Monthly)", level=1)
infra_table = doc.add_table(rows=1, cols=3)
hdr_cells = infra_table.rows[0].cells
hdr_cells[0].text = "Item"
hdr_cells[1].text = "Monthly Cost"
hdr_cells[2].text = "Notes"

infra_items = [
    ("EC2 (App Servers)", "£60–£100", "For 1–2 small servers (auto-scalable)"),
    ("RDS (PostgreSQL DB)", "£50–£80", "Managed, encrypted"),
    ("S3 (File Storage)", "£10–£20", "For reports, logs"),
    ("Redis (optional)", "£15–£30", "Job queue/session cache"),
    ("Load Balancer + Domain", "£25–£40", "AWS ELB or Azure Front Door"),
    ("SSL Certificates", "Free", "Let’s Encrypt or included"),
]
for item in infra_items:
    row_cells = infra_table.add_row().cells
    for i in range(3):
        row_cells[i].text = item[i]

doc.add_paragraph(
    "Estimated Monthly Infrastructure: £160 – £250\nEstimated Annual Infrastructure: £2,000 – £3,000"
)

# Third-party Services
doc.add_heading("3. Third-Party Services (Monthly)", level=1)
third_party_table = doc.add_table(rows=1, cols=3)
hdr_cells = third_party_table.rows[0].cells
hdr_cells[0].text = "Service"
hdr_cells[1].text = "Monthly Cost"
hdr_cells[2].text = "Notes"

third_party_items = [
    ("Twilio SMS", "£40–£100", "£0.04–£0.08/SMS"),
    ("SendGrid (Email)", "Free – £15", "Email delivery"),
    ("Auth0 (Auth)", "Free – £50", "Free up to 7,000 users"),
    ("Sentry / Monitoring", "Free – £30", "Error tracking"),
    ("Pingdom / Uptime", "£10–£20", "Optional"),
]
for item in third_party_items:
    row_cells = third_party_table.add_row().cells
    for i in range(3):
        row_cells[i].text = item[i]

doc.add_paragraph(
    "Estimated Monthly Third-party Services: £100 – £200\nEstimated Annual Services: £1,200 – £2,400"
)

# One-time Costs
doc.add_heading("4. One-Time Tools & Licensing", level=1)
doc.add_paragraph(
    "Most tools and libraries used (e.g. React, Chart.js, jsPDF) are open-source or free."
)
doc.add_paragraph("Estimated One-Time Tooling Cost: £0 – £500")

# Summary
doc.add_heading("5. Total Estimated Budget (Year 1)", level=1)
summary_table = doc.add_table(rows=1, cols=2)
hdr_cells = summary_table.rows[0].cells
hdr_cells[0].text = "Category"
hdr_cells[1].text = "Estimate"

summary_items = [
    ("Development (MVP Build)", "£45,000 – £55,000"),
    ("Infrastructure (Annual)", "£2,000 – £3,000"),
    ("Third-party Services (Annual)", "£1,200 – £2,400"),
    ("One-Time Tools/Licensing", "£0 – £500"),
]
for item in summary_items:
    row_cells = summary_table.add_row().cells
    row_cells[0].text = item[0]
    row_cells[1].text = item[1]

doc.add_paragraph("\nTotal Estimated Year 1 Cost: £48,000 – £61,000")

# Save the document
doc.save("Hopewell_Student_Dashboard_Budget_Estimate.docx")
